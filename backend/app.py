from flask import Flask, request, jsonify
from flask_cors import CORS
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from openai import OpenAI
from dotenv import load_dotenv
import os
import traceback
import re

# Cargar variables de entorno
load_dotenv()

app = Flask(__name__)
CORS(app)

# Inicializar cliente de OpenAI
openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

DATA_PATH = "data/docs"
INDEX_PATH = "data/nicsp_index"

# Usar embeddings de OpenAI para mejor calidad
embeddings = OpenAIEmbeddings(
    model="text-embedding-3-small",
    openai_api_key=os.getenv("OPENAI_API_KEY")
)

def read_pdf(path):
    try:
        import fitz
        doc = fitz.open(path)
        return "\n".join(page.get_text() for page in doc)
    except Exception as e:
        print(f"No se pudo leer PDF {path}: {e}")
        return ""

def read_docx(path):
    try:
        import docx
        d = docx.Document(path)
        return "\n".join(p.text for p in d.paragraphs)
    except Exception as e:
        print(f"No se pudo leer DOCX {path}: {e}")
        return ""

def read_text_like(path):
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            with open(path, encoding=enc) as f:
                return f.read()
        except Exception:
            continue
    print(f"No se pudo decodificar {path}; omitiendo.")
    return ""

def read_generic(path):
    try:
        with open(path, "rb") as f:
            data = f.read()
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                return data.decode(enc)
            except Exception:
                continue
    except Exception as e:
        print(f"No se pudo leer {path}: {e}")
    return ""

# Función auxiliar para limpiar texto
def clean_text(text):
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

# Construir o cargar índice
if not os.path.exists(INDEX_PATH):
    print("Generando base de conocimiento FAISS...")
    docs = []
    
    def read_text_file(path):
        _, ext = os.path.splitext(path)
        ext = ext.lower()
        if ext == ".pdf":
            return read_pdf(path)
        if ext == ".docx":
            return read_docx(path)
        if ext in (".txt", ".md", ".csv", ".json"):
            return read_text_like(path)
        return read_generic(path)
    
    for file in os.listdir(DATA_PATH):
        file_path = os.path.join(DATA_PATH, file)
        if os.path.isdir(file_path):
            continue
        
        text = clean_text(read_text_file(file_path))
        if not text:
            print(f"Omitiendo archivo vacío: {file}")
            continue
        
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=750,             # Tamaño óptimo para listas largas
            chunk_overlap=120,          # Mantener overlap proporcional
            separators=["\n\n", "\n", ". ", ".", " "]
        )
        chunks = splitter.split_text(text)
        for chunk in chunks:
            docs.append(Document(page_content=chunk, metadata={"source": file}))
    
    print(f"✓ Total documentos cargados: {len(docs)}")
    db = FAISS.from_documents(docs, embeddings)
    db.save_local(INDEX_PATH)
else:
    print("Cargando índice FAISS existente...")
    db = FAISS.load_local(INDEX_PATH, embeddings, allow_dangerous_deserialization=True)

# Configuración de OpenAI optimizada
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
OPENAI_MAX_TOKENS = int(os.getenv("OPENAI_MAX_TOKENS", "150"))  # Reducido para respuestas más concisas
OPENAI_TEMPERATURE = float(os.getenv("OPENAI_TEMPERATURE", "0.0"))  # 0.0 para máxima determinismo
OPENAI_TOP_P = float(os.getenv("OPENAI_TOP_P", "0.2"))  # Reducido para mayor enfoque
MAX_CONTEXT_LENGTH = 1000  # Límite de contexto en caracteres

# Almacenar archivos subidos temporalmente
UPLOAD_FOLDER = "data/uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


def call_openai_api(prompt: str) -> str:
    """
    Llama a la API de OpenAI con parámetros optimizados.
    """
    try:
        # Sistema simplificado
        system_prompt = "Experto NICSP. Responde con información del contexto proporcionado."
        
        response = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt}
            ],
            max_tokens=OPENAI_MAX_TOKENS,
            temperature=OPENAI_TEMPERATURE,
            top_p=OPENAI_TOP_P
        )
        
        return response.choices[0].message.content.strip()
    
    except Exception as e:
        error_str = str(e)
        print(f"[ERROR] OpenAI API: {e}")
        
        # Manejo específico de errores comunes
        if "insufficient_quota" in error_str or "429" in error_str:
            return "❌ Cuota de API de OpenAI excedida. Por favor, verifica tu plan de facturación en platform.openai.com/account/billing"
        elif "invalid_api_key" in error_str or "401" in error_str:
            return "❌ API key de OpenAI inválida. Verifica tu configuración."
        
        return f"Error al llamar a OpenAI: {error_str}"


def clean_and_validate_response(response: str, context: str) -> str:
    suspicious = ["remetente", "confirmante", "según mi conocimiento"]
    response_lower = response.lower()
    context_lower = context.lower()
    
    for term in suspicious:
        if term in response_lower and term not in context_lower:
            return "No tengo información suficiente en mis documentos NICSP."
    
    response = response.strip()
    
    #  Permitir respuestas de "no sé" ANTES de validar longitud
    no_info_phrases = ["no encuentro", "no tengo", "no está", "no se encuentra", "no hay información"]
    if any(phrase in response_lower for phrase in no_info_phrases):
        return response
    
    #  Reducir umbral mínimo a 20 caracteres
    if len(response) < 20:
        return "No encuentro información específica sobre eso en los documentos."
    
    return response


def build_context_from_docs(docs):
    """Construir contexto limitado desde documentos MMR."""
    if not docs:
        return "No hay información relevante."
    
    context_parts = []
    total_chars = 0
    
    for doc in docs:
        content = doc.page_content.strip()
        
        # Limitar cada chunk a ~300 chars
        if len(content) > 300:
            content = content[:300] + "..."
        
        if total_chars + len(content) > MAX_CONTEXT_LENGTH:
            break
            
        context_parts.append(content)
        total_chars += len(content)
    
    context = "\n\n".join(context_parts)
    
    # Garantizar que no exceda el límite
    if len(context) > MAX_CONTEXT_LENGTH:
        context = context[:MAX_CONTEXT_LENGTH] + "..."
    
    return context


@app.post("/chat")
def chat():
    message = request.json.get("message", "")
    if not message:
        return jsonify({"error": "No message provided"}), 400
    
    try:
        # Usar MMR (Maximum Marginal Relevance) para mejor diversidad
        docs = db.max_marginal_relevance_search(
            message, 
            k=3,  # Reducido a 3 documentos
            fetch_k=10  # Buscar en 10, seleccionar 3 más diversos
        )
        
        context = build_context_from_docs(docs)
        
        print(f"\n{'='*60}")
        print(f"PREGUNTA: {message}")
        print(f"DOCS MMR: {len(docs)}")
        print(f"CONTEXTO ({len(context)} chars):\n{context}")
        print(f"{'='*60}\n")
        
        # Prompt simplificado
        prompt = f"""Contexto:
{context}

Pregunta: {message}

Responde solo con el contexto. Máximo 3 oraciones."""
        
        answer = call_openai_api(prompt)
        
        if not answer or answer.strip() == "":
            return jsonify({"error": "No se pudo generar respuesta"}), 500
        
        answer = clean_and_validate_response(answer, context)
        
        print(f"[RESPUESTA]: {answer}\n")
        
        return jsonify({"response": answer})
    
    except Exception as e:
        error_msg = str(e)
        print("[ERROR]:", traceback.format_exc())
        
        # Mensajes de error más específicos
        if "insufficient_quota" in error_msg or "429" in error_msg:
            return jsonify({
                "error": "⚠️ Cuota de OpenAI excedida. Agrega créditos en platform.openai.com/account/billing"
            }), 429
        elif "invalid_api_key" in error_msg or "401" in error_msg:
            return jsonify({
                "error": "❌ API key de OpenAI inválida. Verifica la configuración del servidor."
            }), 401
        
        return jsonify({"error": "Error interno del servidor"}), 500


@app.post("/upload")
def upload_file():
    """
    Endpoint para subir archivos PDF y procesarlos con OpenAI.
    """
    if 'file' not in request.files:
        return jsonify({"error": "No se proporcionó ningún archivo"}), 400
    
    file = request.files['file']
    question = request.form.get('question', '')
    
    if file.filename == '':
        return jsonify({"error": "Nombre de archivo vacío"}), 400
    
    if not file.filename.endswith('.pdf'):
        return jsonify({"error": "Solo se permiten archivos PDF"}), 400
    
    try:
        # Guardar archivo temporalmente
        file_path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(file_path)
        
        # Subir archivo a OpenAI
        with open(file_path, "rb") as f:
            uploaded_file = openai_client.files.create(
                file=f,
                purpose="assistants"
            )
        
        # Procesar pregunta con el archivo (simplificado)
        prompt = f"""Documento adjunto sobre NICSP.

Pregunta: {question}

Responde solo con información del documento. Máximo 3 oraciones."""
        
        response = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": "Experto NICSP. Responde solo con información del documento."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            max_tokens=OPENAI_MAX_TOKENS,
            temperature=OPENAI_TEMPERATURE,
            top_p=OPENAI_TOP_P
        )
        
        answer = response.choices[0].message.content.strip()
        
        # Limpiar archivo temporal
        os.remove(file_path)
        
        # Eliminar archivo de OpenAI después de procesarlo
        try:
            openai_client.files.delete(uploaded_file.id)
        except Exception as delete_error:
            print(f"[WARNING] No se pudo eliminar archivo de OpenAI: {delete_error}")
        
        print(f"[UPLOAD] Archivo: {file.filename}, Pregunta: {question}")
        print(f"[RESPUESTA]: {answer}\n")
        
        return jsonify({
            "response": answer,
            "filename": file.filename
        })
    
    except Exception as e:
        print(f"[ERROR] Upload: {traceback.format_exc()}")
        try:
            file_path = os.path.join(UPLOAD_FOLDER, file.filename)
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as cleanup_error:
            print(f"[WARNING] No se pudo limpiar archivo: {cleanup_error}")
        return jsonify({"error": f"Error al procesar el archivo: {str(e)}"}), 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)